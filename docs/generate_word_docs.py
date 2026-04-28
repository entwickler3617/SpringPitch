from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

ROOT = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = ROOT

INSTALL_DOC = os.path.join(OUT_DIR, 'SpringPitch_설치가이드.docx')
USER_DOC = os.path.join(OUT_DIR, 'SpringCalculator_사용자매뉴얼.docx')

PROJECT_NAME = 'SpringPitch / SpringCalculator'

SYSTEM_REQ = [
    '운영체제: Windows 10/11 (64-bit)',
    '권장 사양: 8GB RAM 이상, 2GB 여유 디스크 공간',
    'Python 3.13 (개발/빌드 시), 런타임은 EXE 단독 실행 가능',
]

DEPENDENCIES = [
    'numpy',
    'scipy',
    'pandas',
    'matplotlib (Agg 백엔드 사용)',
    'openpyxl',
    'pyinstaller (빌드용)',
]

ENV_VARS = [
    'SPRING_CREATE_ZERO1_IF_MISSING=1   # zero-1 시트가 없을 때만 생성 허용',
    'SPRING_THETA_MODE=start0_unwrapped # 시작 각도 0 기준 언랩',
    'SPRING_THETA_POSITIVE=true         # 각도 양의 증가 방향 강제',
    'SPRING_START_STRATEGY=min_radius   # 시작점 선택 전략',
]

BUILD_CMD = (
    'pyinstaller --noconfirm --onefile --name SpringCalculator f:/SpringPitch/detailed_center_analysis.py'
)

RUN_TIPS = [
    '동일 폴더에 다음 파일을 함께 둡니다: SpringCalculator.exe, TK1.xlsx, 대상 .igs 파일 1개',
    'Excel에서 TK1.xlsx가 열려 있지 않도록 합니다 (쓰기 충돌 방지)',
    '실행 후 zero-1 시트 2행부터 결과가 추가되며 서식은 변경되지 않습니다',
    '이미지: spring_detailed_analysis.png, new_result_report.png 생성',
]

TROUBLESHOOT = [
    '[TK1.xlsx 없음] TK1.xlsx가 실행 폴더에 반드시 존재해야 합니다',
    "[시트 없음] 기본은 zero-1 시트를 요구합니다. 시트가 없다면 환경변수 'SPRING_CREATE_ZERO1_IF_MISSING=1' 설정 후 실행 (필요 시에만)",
    '[파일 사용 중] Excel에서 TK1.xlsx가 열려 있으면 저장 실패. 닫은 뒤 재실행',
    '[IGS 미발견] 실행 폴더에 *.igs 파일이 1개 이상 존재해야 합니다',
]

COMPONENTS = [
    'detailed_center_analysis.py: 메인 분석/출력 엔진 (EXE 엔트리포인트)',
    'iges_parser.py: IGS 포인트 파서(섹션/정규식 보조)',
    'plot/분석 루틴: 국소좌표 변환, 언랩된 각도, 피치/반경 등 계산',
]


def set_korean_style(document):
    try:
        style = document.styles['Normal']
        font = style.font
        font.name = 'Malgun Gothic'
        font.size = Pt(11)
        rFonts = style.element.rPr.rFonts
        rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    except Exception:
        pass


def add_heading(document, text: str, level: int = 1):
    h = document.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_bullets(document, items):
    for it in items:
        p = document.add_paragraph(it, style='List Bullet')
        p.paragraph_format.space_after = Pt(0)


def add_code_block(document, lines):
    for ln in lines if isinstance(lines, list) else [lines]:
        p = document.add_paragraph()
        run = p.add_run(ln)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)


def build_install_guide(path: str):
    doc = Document()
    set_korean_style(doc)

    add_heading(doc, f'{PROJECT_NAME} 설치 가이드', 0)
    doc.add_paragraph('본 문서는 개발환경 준비, 라이브러리 설치, 빌드 및 실행에 관한 절차를 정리합니다.')

    add_heading(doc, '1. 시스템 요구사항', 1)
    add_bullets(doc, SYSTEM_REQ)

    add_heading(doc, '2. 개발환경 설치 및 설정', 1)
    add_bullets(doc, [
        'Python 3.13 설치 (Microsoft Store 또는 python.org)',
        '프로젝트 루트에서 가상환경 생성 및 활성화',
    ])
    add_code_block(doc, [
        'PowerShell:',
        'python -m venv .venv',
        '& .\\.venv\\Scripts\\Activate.ps1',
    ])

    add_heading(doc, '3. 필요 라이브러리/모듈 설치', 1)
    add_bullets(doc, DEPENDENCIES)
    add_code_block(doc, ['pip install numpy scipy pandas matplotlib openpyxl pyinstaller'])

    add_heading(doc, '4. 프로그램 구성요소', 1)
    add_bullets(doc, COMPONENTS)

    add_heading(doc, '5. 빌드(응용프로그램 생성) 방법', 1)
    doc.add_paragraph('PyInstaller를 이용하여 단일 실행파일(Onefile) EXE를 생성합니다:')
    add_code_block(doc, [BUILD_CMD])
    add_bullets(doc, [
        '성공 시 dist/SpringCalculator.exe 생성',
        '루트로 복사하여 배포용으로 사용 가능',
    ])

    add_heading(doc, '6. 실행 방법(개발 환경)', 1)
    add_bullets(doc, [
        '분석용 .igs 및 TK1.xlsx 파일을 실행 경로에 준비',
        'Python으로 직접 실행 시 detailed_center_analysis.py를 호출',
    ])
    add_code_block(doc, ['python f:/SpringPitch/detailed_center_analysis.py'])

    add_heading(doc, '7. 환경 변수(선택적)', 1)
    add_bullets(doc, ENV_VARS)

    add_heading(doc, '8. 문제 해결 가이드', 1)
    add_bullets(doc, TROUBLESHOOT)

    doc.save(path)


def build_user_manual(path: str):
    doc = Document()
    set_korean_style(doc)

    add_heading(doc, 'SpringCalculator.exe 사용자 매뉴얼', 0)
    doc.add_paragraph('본 문서는 SpringCalculator.exe의 생성 및 실행 방법을 안내합니다.')

    add_heading(doc, '1. 응용프로그램 생성방법', 1)
    add_bullets(doc, [
        '개발환경 준비: Python 3.13, 가상환경, 필수 라이브러리 설치',
        'PyInstaller로 빌드 (onefile, 이름: SpringCalculator)',
    ])
    add_code_block(doc, [BUILD_CMD])

    add_heading(doc, '2. 응용프로그램 실행방법', 1)
    add_bullets(doc, [
        'SpringCalculator.exe, TK1.xlsx, 대상 .igs 파일을 동일 폴더에 둠',
        'TK1.xlsx는 반드시 닫힌 상태에서 실행',
        '실행 후 zero-1 시트 2행부터 결과 데이터가 추가',
        '시트 서식/폭 등은 변경하지 않음',
        'spring_detailed_analysis.png, new_result_report.png 생성',
    ])

    add_heading(doc, '3. 입력/출력 규칙', 1)
    add_bullets(doc, [
        '*.igs: 실행 폴더에서 자동으로 첫 번째 파일 사용',
        'TK1.xlsx: zero-1 시트에만 데이터 추가, 새로운 파일/시트 생성 금지(기본)',
        '필요 시 환경변수로 zero-1 시트 생성 허용 (SPRING_CREATE_ZERO1_IF_MISSING=1)',
    ])

    add_heading(doc, '4. 실행 전 체크리스트', 1)
    add_bullets(doc, RUN_TIPS)

    add_heading(doc, '5. 자주 묻는 질문(FAQ) / 문제 해결', 1)
    add_bullets(doc, TROUBLESHOOT)

    add_heading(doc, '6. 버전/구성요소 정보', 1)
    add_bullets(doc, [
        'Matplotlib 백엔드: Agg (GUI 의존성 제거)',
        'Excel 기록: zero-1 시트 2행부터, 번호(No)는 1부터 증가',
        '중복 저장/경고 제거 로직 반영',
    ])

    doc.save(path)


if __name__ == '__main__':
    os.makedirs(OUT_DIR, exist_ok=True)
    build_install_guide(INSTALL_DOC)
    build_user_manual(USER_DOC)
    print('Generated:')
    print(' -', INSTALL_DOC)
    print(' -', USER_DOC)
