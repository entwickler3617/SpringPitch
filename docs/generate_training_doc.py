#!/usr/bin/env python3
"""
스프링 형상추출 프로그램 교육자료 생성 스크립트
- 하루 분량 교육자료를 DOCX 형식으로 생성
- 대상: 전문개발자가 아닌 엔지니어/기술자
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_PATH = os.path.join(ROOT, 'docs', '스프링형상추출_교육자료.docx')


# ── 유틸리티 함수 ──────────────────────────────────────────────

def set_korean_style(doc):
    """맑은 고딕 기본 스타일 설정"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(11)
    try:
        rFonts = style.element.rPr.rFonts
        rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    except Exception:
        pass
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def add_para(doc, text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_bullets(doc, items, indent_level=0):
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        if indent_level > 0:
            p.paragraph_format.left_indent = Cm(indent_level * 0.6)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.space_after = Pt(2)


def add_code(doc, code_text, language='python'):
    """코드 블록을 회색 배경으로 표시"""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x20, 0x20, 0x60)


def add_tip_box(doc, title, content):
    """팁/참고 박스 (표 형태)"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    # 셀 배경색
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'E8F4FD')
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)
    p1 = cell.paragraphs[0]
    run1 = p1.add_run(f'💡 {title}')
    run1.bold = True
    run1.font.size = Pt(10)
    p2 = cell.add_paragraph(content)
    p2.runs[0].font.size = Pt(10) if p2.runs else None
    doc.add_paragraph()  # 간격


def add_simple_table(doc, headers, rows, col_widths=None):
    """간단한 표 추가"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    # 데이터
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    doc.add_paragraph()  # 간격
    return table


# ── 본문 작성 ──────────────────────────────────────────────────

def build_training_doc():
    doc = Document()
    set_korean_style(doc)

    # ━━━━━━━━ 표지 ━━━━━━━━
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading('스프링 형상추출 프로그램\n파이썬 교육자료', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— 하루 교육과정 —')
    run.font.size = Pt(14)
    run.bold = True
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('대상: 엔지니어 / 기술자 (비전문 개발자)\n').font.size = Pt(11)
    meta.add_run('프로젝트: SpringPitch (스프링 IGS 형상 데이터 분석)\n').font.size = Pt(11)
    meta.add_run('작성일: 2026-02-09').font.size = Pt(11)

    doc.add_page_break()

    # ━━━━━━━━ 목차 ━━━━━━━━
    add_heading(doc, '목차', level=1)
    toc_items = [
        '1. 파이썬 개념 소개',
        '   A. 기본 문법 (데이터, 변수, 함수)',
        '   B. 형상추출 프로그램 이해에 필요한 파이썬 장점',
        '   C. 본 프로젝트에서 사용하는 주요 라이브러리',
        '2. 프로그램 소스 코드 소개',
        '   A. 알고리즘: IGES 형상 데이터 추출 과정(프로세스) 이해',
        '   B. 플로차트: 프로그램 소스 구성 및 데이터 교환',
        '   C. 소스 확보/디버깅 수정 방법: AI 활용',
        '   D. 프로그램 실행 방법',
        '   E. 프로그램 출력물 이해',
        '3. 기타',
        '   A. 자주 묻는 질문 (Q&A)',
        '   B. 참고 자료 및 용어 정리',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(1)

    doc.add_page_break()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 1. 파이썬 개념 소개
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_heading(doc, '1. 파이썬(Python) 개념 소개', level=1)

    add_para(doc, '파이썬은 1991년 귀도 반 로섬이 만든 프로그래밍 언어로, '
             '"읽기 쉽고 배우기 쉬운" 것을 철학으로 삼고 있습니다. '
             '전 세계에서 가장 많이 사용되는 언어 중 하나이며, '
             '특히 데이터 분석, 과학 계산, 자동화 분야에서 널리 쓰입니다.')

    # ── 1A. 기본 문법 ──
    add_heading(doc, 'A. 기본 문법 (데이터, 변수, 함수)', level=2)

    # 데이터 타입
    add_heading(doc, '① 데이터 타입', level=3)
    add_para(doc, '파이썬에서 다루는 기본 데이터 종류입니다:')
    add_simple_table(doc,
        ['데이터 타입', '설명', '예시'],
        [
            ['정수 (int)', '소수점 없는 숫자', '42, -7, 0'],
            ['실수 (float)', '소수점이 있는 숫자', '3.14, -0.001, 1.0'],
            ['문자열 (str)', '텍스트 데이터', '"Hello", \'스프링\''],
            ['리스트 (list)', '여러 값을 순서대로 모음', '[1, 2, 3], ["a", "b"]'],
            ['딕셔너리 (dict)', '이름표(키)로 값을 관리', '{"높이": 128.9, "반경": 50.0}'],
            ['불리언 (bool)', '참/거짓', 'True, False'],
        ])

    add_para(doc, '본 프로젝트에서는 주로 실수(float)와 리스트를 사용하여 '
             '3차원 좌표 데이터를 처리합니다.', italic=True)

    # 변수
    add_heading(doc, '② 변수', level=3)
    add_para(doc, '변수는 데이터를 담아두는 이름표입니다. '
             '파이썬에서는 별도의 "타입 선언" 없이 바로 값을 할당할 수 있습니다:')
    add_code(doc, '''
# 스프링 포인트 좌표를 변수에 저장
x = -284.949      # X 좌표 (실수)
y = 98.352         # Y 좌표
z = 383.726        # Z 좌표

# 여러 좌표를 리스트로 묶기
point = [x, y, z]                 # 하나의 3D 점
points = [[1,2,3], [4,5,6]]      # 점 여러 개

# 분석 결과를 딕셔너리로 정리
result = {
    "스프링_높이": 128.869,
    "평균_반지름": 49.955,
    "총_회전수": 7.2
}
''')

    # 함수
    add_heading(doc, '③ 함수', level=3)
    add_para(doc, '함수는 특정 작업을 묶어서 이름을 붙인 것입니다. '
             '"def" 키워드로 정의하고, 필요할 때 이름을 불러 사용합니다:')
    add_code(doc, '''
# 함수 정의 예시: 두 점 사이의 거리 계산
def distance(p1, p2):
    """두 3D 점 사이의 유클리드 거리를 계산"""
    dx = p1[0] - p2[0]
    dy = p1[1] - p2[1]
    dz = p1[2] - p2[2]
    return (dx**2 + dy**2 + dz**2) ** 0.5

# 함수 호출
d = distance([0, 0, 0], [3, 4, 0])
print(d)   # 결과: 5.0
''')

    add_para(doc, '본 프로젝트의 핵심 함수들:')
    add_simple_table(doc,
        ['함수명', '파일', '역할'],
        [
            ['parse_igs_points()', 'iges_parser.py', 'IGS 파일에서 3D 좌표를 추출'],
            ['find_spring_axis_pca()', 'detailed_center_analysis.py', 'PCA로 스프링 주축 계산'],
            ['to_local_coordinates()', 'detailed_center_analysis.py', '글로벌→로컬 좌표 변환'],
            ['resample_curve()', 'detailed_center_analysis.py', '1000개 등간격 리샘플링'],
            ['write_excel_tk1()', 'detailed_center_analysis.py', '결과를 Excel에 기록'],
        ])

    # 제어문
    add_heading(doc, '④ 조건문과 반복문', level=3)
    add_para(doc, '프로그램 흐름을 제어하는 기본 구조입니다:')
    add_code(doc, '''
# 조건문: 파일이 존재하는지 확인
if os.path.exists("TK1.xlsx"):
    print("엑셀 파일 발견!")
else:
    print("파일이 없습니다.")

# 반복문: 모든 포인트를 순회하며 처리
for point in points:
    x, y, z = point[0], point[1], point[2]
    r = (x**2 + y**2) ** 0.5   # 반지름 계산
''')

    add_tip_box(doc, '참고: 파이썬의 들여쓰기',
                '파이썬은 중괄호({}) 대신 들여쓰기(스페이스 4칸)로 코드 블록을 구분합니다. '
                'if, for, def 등 뒤에 오는 코드는 반드시 들여쓰기가 필요합니다.')

    # ── 1B. 형상추출 프로그램을 위한 파이썬 장점 ──
    add_heading(doc, 'B. 형상추출 프로그램 이해에 필요한 파이썬 장점', level=2)

    add_numbered(doc, [
        '풍부한 수학/과학 라이브러리: NumPy(행렬 계산), SciPy(스플라인 보간), '
        'Matplotlib(그래프) 등을 한 줄로 불러 쓸 수 있습니다.',
        '읽기 쉬운 문법: C/C++에 비해 코드가 직관적이어서, '
        '엔지니어도 알고리즘 의도를 빠르게 파악할 수 있습니다.',
        '빠른 프로토타이핑: 아이디어를 즉시 코드로 구현하고 결과를 확인할 수 있어, '
        'CAD 데이터 분석 실험에 적합합니다.',
        '크로스플랫폼: Windows, Mac, Linux 어디서든 동일하게 실행됩니다.',
        '단독 실행파일 생성: PyInstaller로 .exe를 만들면 Python 없이도 배포 가능합니다.',
        'AI 도구와의 친화성: ChatGPT, Copilot 등 AI 코딩 도우미가 파이썬을 가장 잘 지원합니다.',
    ])

    # ── 1C. 주요 라이브러리 ──
    add_heading(doc, 'C. 본 프로젝트에서 사용하는 주요 라이브러리', level=2)

    add_simple_table(doc,
        ['라이브러리', '용도', '프로젝트 내 활용'],
        [
            ['numpy', '수치 계산/행렬 연산', '좌표 배열 관리, PCA 고유값 분해, 벡터 연산 전반'],
            ['matplotlib', '2D/3D 그래프 시각화', '스프링 3D 플롯, R/θ/Z 그래프, 분석 보고서 차트'],
            ['openpyxl', 'Excel 파일 읽기/쓰기', 'TK1.xlsx에 분석 결과 기록 (서식 보존)'],
            ['pandas', '표 형태 데이터 처리', 'Excel 시트 읽기, 데이터프레임 편의 기능'],
            ['scipy', '과학 계산 (스플라인 등)', 'B-Spline 보간을 이용한 곡선 표준화'],
            ['re (내장)', '정규표현식 (텍스트 패턴 매칭)', 'IGS 파일에서 좌표 매칭 (폴백 파서)'],
            ['math (내장)', '수학 함수', '삼각함수, 각도 변환 등'],
        ])

    add_code(doc, '''
# 라이브러리 불러오기 예시 (본 프로젝트 코드에서 발췌)
import numpy as np               # 수치 계산
import matplotlib.pyplot as plt   # 그래프
from openpyxl import load_workbook  # Excel
import re                         # 정규식
''')

    doc.add_page_break()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 2. 프로그램 소스 코드 소개
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_heading(doc, '2. 프로그램 소스 코드 소개', level=1)

    # ── 2A. 알고리즘 ──
    add_heading(doc, 'A. 알고리즘: IGES 형상 데이터 추출 과정', level=2)

    add_heading(doc, '① IGES 파일이란?', level=3)
    add_para(doc, 'IGES(Initial Graphics Exchange Specification)는 CAD 시스템 간에 '
             '기하학적 데이터를 교환하기 위한 국제 표준 포맷입니다. '
             '확장자는 .igs 또는 .iges이며, ASCII 텍스트 기반입니다.')

    add_para(doc, 'IGES 파일은 80자 고정 길이 줄로 구성되며, 5개 섹션으로 나뉩니다:', bold=True)
    add_simple_table(doc,
        ['섹션', '식별 문자', '역할'],
        [
            ['Start (S)', 'S', '파일 설명 (자유 텍스트)'],
            ['Global (G)', 'G', '메타데이터: 구분자, 단위(mm/inch), 파일 정보'],
            ['Directory (D)', 'D', '엔터티 목록: 종류, 위치, 변환행렬 포인터 등'],
            ['Parameter (P)', 'P', '실제 기하 데이터: 좌표값, 곡선 파라미터 등'],
            ['Terminate (T)', 'T', '파일 끝 표시 및 각 섹션 줄 수 요약'],
        ])

    add_para(doc, '본 프로젝트에서 중요한 엔터티:', bold=True)
    add_simple_table(doc,
        ['엔터티 번호', '이름', '설명'],
        [
            ['116', 'Point', '3D 포인트 (X, Y, Z 좌표)'],
            ['124', 'Transformation Matrix', '좌표 변환 행렬 (3×4 회전+이동)'],
            ['110', 'Line', '직선 (시작점, 끝점)'],
            ['126', 'Rational B-Spline Curve', 'NURBS 곡선'],
        ])

    add_heading(doc, '② 전체 분석 프로세스 (5단계)', level=3)
    add_para(doc, '스프링 형상추출은 아래 5단계로 진행됩니다:')

    add_para(doc, 'Step 1: IGS 포인트 로드', bold=True)
    add_bullets(doc, [
        'IGS 파일에서 엔터티 116(Point)의 X, Y, Z 좌표를 추출',
        '124(Transformation Matrix) 엔터티가 참조되면 좌표 변환 적용',
        '단위를 mm로 통일 (Global 섹션에서 단위 정보 확인)',
        '예시: TK1_FRT_zero-1_251014.igs → 3,141개 포인트 추출',
    ])

    add_para(doc, 'Step 2: 좌표 정규화 (1,000개 등간격 표준화)', bold=True)
    add_bullets(doc, [
        '3,141개 원시 포인트를 스프링 경로를 따라 정렬 (나선 파라미터 기반)',
        'PCA(주성분분석)로 스프링 주축(중심선) 방향을 찾음',
        '국소 좌표계로 변환: 주축 = 로컬 X축',
        '호 길이(arc-length) 기반으로 정확히 1,000개 등간격 포인트로 리샘플링',
        '결과: (x_norm, y_norm, z_norm) = 표준화된 국소 좌표',
    ])

    add_para(doc, 'Step 3: 원통 좌표 변환 (xyz → Rθz)', bold=True)
    add_bullets(doc, [
        '국소 직교좌표 (x_norm, y_norm, z_norm) → 원통좌표 (R, θ, Z)',
        'R = √(y_norm² + z_norm²) : 축으로부터의 거리 (반지름)',
        'θ = atan2(z_norm, y_norm) : 축 주위의 각도',
        'Z = x_norm : 축 방향 위치 (스프링 높이)',
        'θ는 unwrap하여 연속값으로 처리 후, 시작점에서 0°로 설정',
    ])

    add_para(doc, 'Step 4: 기초 치수 산출', bold=True)
    add_bullets(doc, [
        'Turn (회전수) = θ_unwrapped / 360°',
        'Diameter (수직) = 2 × R',
        'Diameter (벡터) = 반 바퀴 앞의 점과의 직선 거리',
        'Height (높이) = Z값의 변화',
        'Pitch (피치) = 1회전당 축방향 이동 거리',
        'min_Pitch = 전체 구간에서의 최소 피치값',
    ])

    add_para(doc, 'Step 5: 3점 호 보간 정밀 치수 (현재 비활성)', bold=True)
    add_bullets(doc, [
        '세 점으로 원호를 피팅하여 정밀 반지름/피치 계산 (옵션)',
        '현재 요구사항에 따라 생략 상태',
    ])

    add_heading(doc, '③ PCA(주성분분석)란?', level=3)
    add_para(doc, 'PCA는 데이터의 분산이 가장 큰 방향(주축)을 찾는 통계적 방법입니다.')
    add_bullets(doc, [
        '스프링 포인트는 나선형으로 분포 → 가장 분산이 큰 방향 = 스프링 축',
        '본 데이터: 주축이 전체 분산의 88.7%를 설명 (매우 뚜렷한 나선)',
        '나머지 두 축(5.7%, 5.6%)은 원형 단면에 해당',
    ])

    add_code(doc, '''
# PCA 코드 핵심 (detailed_center_analysis.py에서 발췌)
def find_spring_axis_pca(points):
    center = np.mean(points, axis=0)        # 무게 중심
    centered = points - center               # 중심 이동
    cov_matrix = np.cov(centered.T)          # 공분산 행렬
    eigenvalues, eigenvectors = np.linalg.eigh(cov_matrix)
    # 고유값이 큰 순서로 정렬 → 첫 번째가 주축
    idx = np.argsort(eigenvalues)[::-1]
    principal_axis = eigenvectors[:, idx[0]]
    return center, principal_axis
''')

    add_heading(doc, '④ 원통좌표계 (R, θ, Z) 변환 원리', level=3)
    add_para(doc, '직교좌표 (x, y, z)를 원통좌표 (R, θ, Z)로 변환하는 것은 '
             '스프링의 나선 구조를 효과적으로 표현하기 위함입니다:')

    add_simple_table(doc,
        ['원통좌표', '공식', '물리적 의미'],
        [
            ['R (반지름)', '√(y² + z²)', '축으로부터의 거리 → 코일 반경'],
            ['θ (각도)', 'atan2(z, y)', '축 주위의 회전 각도 → 회전 위치'],
            ['Z (높이)', 'x (축방향)', '축을 따른 위치 → 스프링 높이'],
        ])

    doc.add_page_break()

    # ── 2B. 플로차트 ──
    add_heading(doc, 'B. 플로차트: 프로그램 소스 구성 및 데이터 교환', level=2)

    add_heading(doc, '① 파일 구성도', level=3)
    add_para(doc, '프로젝트의 주요 파이썬 파일과 역할:', bold=True)

    add_simple_table(doc,
        ['파일명', '역할', '주요 함수/기능'],
        [
            ['detailed_center_analysis.py', '메인 분석 엔진\n(프로그램 진입점)', 
             'main(), parse_igs_points()\nfind_spring_axis_pca()\nto_local_coordinates()\nresample_curve()\nwrite_excel_tk1()'],
            ['iges_parser.py', 'IGS 파일 전문 파서',
             'parse_iges_points()\n섹션(S/G/D/P/T) 분석\n변환행렬(124) 적용\n단위(mm) 정규화'],
            ['arc_metrics.py', '3점 호 보간 메트릭',
             'circle_from_3pts()\narc_refine_radius_and_pitch()'],
            ['excel_formulas.py', 'Excel 수식 복제 모듈',
             'compute_basic_metrics_from_cyl()\nTurn, Radius, Diameter,\nPitch 계산'],
            ['plot_tk1_excel.py', 'Excel 데이터 시각화',
             'TK1.xlsx 읽어서\nx_norm/y_norm/z_norm 3D 플롯\nR/θ/Z vs index 플롯'],
            ['final_report_generator.py', '종합 보고서 차트 생성',
             '3D 시각화, 메트릭 요약,\n분포 히스토그램'],
            ['report_generator.py', 'Markdown 보고서 생성',
             'generate_report_md()'],
            ['tools/count_igs_points.py', 'IGS 포인트 수 확인 도구',
             'main()'],
            ['tools/export_xyz_and_plot.py', '원시 좌표 CSV 내보내기\n+ 3D 미리보기',
             'main()'],
        ])

    add_heading(doc, '② 데이터 흐름도 (텍스트 다이어그램)', level=3)
    add_code(doc, '''
┌─────────────────────────────────────────────────────────────────┐
│                    프로그램 실행 흐름                             │
├─────────────────────────────────────────────────────────────────┤
│                                                                 │
│  [입력] .igs 파일 (예: TK1_FRT_zero-1_251014.igs)               │
│    │                                                            │
│    ▼                                                            │
│  ┌───────────────────────────────┐                              │
│  │ Step 1: IGS 파싱              │                              │
│  │ (iges_parser.py)              │                              │
│  │ • S/G/D/P/T 섹션 분리          │                              │
│  │ • 엔터티 116(Point) 추출       │                              │
│  │ • 변환행렬(124) 적용           │                              │
│  │ • 단위 mm 정규화               │                              │
│  └──────────┬────────────────────┘                              │
│             │ 원시 좌표 (3,141개)                                 │
│             ▼                                                    │
│  ┌───────────────────────────────┐                              │
│  │ Step 2: 정규화/표준화          │                              │
│  │ (detailed_center_analysis.py) │                              │
│  │ • PCA → 주축 결정              │                              │
│  │ • 나선 파라미터 기반 정렬       │                              │
│  │ • 1,000개 등간격 리샘플링       │                              │
│  │ • 국소좌표 변환                │                              │
│  └──────────┬────────────────────┘                              │
│             │ (x_norm, y_norm, z_norm) × 1,000                  │
│             ▼                                                    │
│  ┌───────────────────────────────┐                              │
│  │ Step 3: 원통좌표 변환          │                              │
│  │ • R = √(y²+z²)               │                              │
│  │ • θ = atan2(z, y)             │                              │
│  │ • Z = x_norm                  │                              │
│  └──────────┬────────────────────┘                              │
│             │ (R, θ, Z) × 1,000                                 │
│             ▼                                                    │
│  ┌───────────────────────────────┐                              │
│  │ Step 4: 기초 치수 산출         │                              │
│  │ (excel_formulas.py 참고)      │                              │
│  │ • Turn, Diameter, Height      │                              │
│  │ • Pitch, min_Pitch            │                              │
│  └──────────┬────────────────────┘                              │
│             │                                                    │
│             ▼                                                    │
│  [출력]                                                         │
│  ├── TK1.xlsx (zero-1 시트: A~X 컬럼)                           │
│  ├── spring_detailed_analysis.png (분석 시각화)                   │
│  ├── new_result_report.png (결과 보고서 차트)                     │
│  └── output/points_xyz.csv (원시 좌표 CSV)                      │
│                                                                 │
└─────────────────────────────────────────────────────────────────┘
''')

    add_heading(doc, '③ Excel 시트(zero-1) 컬럼 구조', level=3)
    add_para(doc, 'TK1.xlsx의 zero-1 시트에 기록되는 핵심 컬럼:', bold=True)
    add_simple_table(doc,
        ['컬럼', '내용', '단위', '설명'],
        [
            ['A (No)', '순번', '-', '1부터 시작'],
            ['B~D (x, y, z)', '원시 좌표', 'mm', 'Step1: 가장 가까운 원본 좌표'],
            ['E~G (x_norm, y_norm, z_norm)', '정규화 좌표', 'mm', 'Step2: 등간격 표준화 좌표'],
            ['H (R)', '반지름', 'mm', 'Step3: 축~포인트 거리'],
            ['I (θ)', '각도', '°(도)', 'Step3: 래핑된 각도 (-180~180)'],
            ['J (Z)', '높이', 'mm', 'Step3: 축방향 위치'],
            ['N (Turn)', '회전수', 'rev', 'Step4: 누적 회전수'],
            ['Q (Perp. Diam.)', '수직 직경', 'mm', 'Step4: 2×R'],
            ['S (Vec. Diam.)', '벡터 직경', 'mm', 'Step4: 대면 점 거리'],
            ['U (Height)', '상대 높이', 'mm', 'Step4: 기준점 대비'],
            ['W (Pitch)', '피치', 'mm', 'Step4: 1회전당 높이'],
            ['X (min_Pitch)', '최소 피치', 'mm', 'Step4: 전체 최소값'],
        ])

    doc.add_page_break()

    # ── 2C. AI 활용 디버깅 ──
    add_heading(doc, 'C. 소스 확보 / 디버깅 수정 방법: AI 활용', level=2)

    add_heading(doc, '① 소스 코드 확보 방법', level=3)
    add_bullets(doc, [
        '프로젝트 폴더(SpringPitch/)에서 모든 .py 파일을 가져옴',
        '원본 코드는 버전 관리 시스템(Git)으로 이력 추적 권장',
        'VS Code(무료 편집기)에서 파일 열기 → 코드 읽기/수정 가능',
    ])

    add_heading(doc, '② AI를 활용한 코드 이해 및 수정', level=3)
    add_para(doc, '전문 개발자가 아니더라도 AI 도구를 활용하면 코드를 이해하고 '
             '수정할 수 있습니다:', bold=True)

    add_para(doc, '(a) ChatGPT / Claude 활용법', bold=True)
    add_bullets(doc, [
        '코드 전체 또는 함수를 복사하여 AI에게 "이 코드가 무엇을 하는지 설명해줘"라고 질문',
        '"이 함수에서 반지름 계산 공식을 R = 2*r로 바꾸려면?" 같은 구체적 수정 요청',
        '에러 메시지를 그대로 붙여넣고 "이 에러의 원인과 해결방법은?" 질문',
    ])

    add_para(doc, '(b) VS Code + GitHub Copilot 활용법', bold=True)
    add_bullets(doc, [
        'VS Code에서 Copilot 확장 설치 → 코드 편집 중 자동 완성 제안',
        'Ctrl+I (인라인 채팅)로 코드 설명 요청 또는 수정 지시',
        '채팅 패널에서 프로젝트 전체에 대한 질문 가능',
    ])

    add_tip_box(doc, 'AI 활용 팁',
                '코드를 질문할 때는 "파일 전체"보다 "특정 함수" 단위로 질문하면 '
                '더 정확한 답변을 받을 수 있습니다.\n\n'
                '예시 프롬프트:\n'
                '"아래 parse_igs_points 함수에서 엔터티 116의 좌표를 추출하는 부분을 설명해줘.\n'
                '특히 정규식 패턴이 어떤 숫자 형식을 매칭하는지 알고 싶어."')

    add_heading(doc, '③ 디버깅(오류 해결) 기본 절차', level=3)
    add_numbered(doc, [
        '에러 메시지 확인: 터미널(명령 프롬프트)에 출력된 오류 내용 읽기',
        'AI에 질문: 에러 메시지 + 관련 코드를 AI에 붙여넣기',
        '수정 적용: AI가 제안한 수정을 코드에 반영',
        '재실행: 수정 후 프로그램을 다시 실행하여 확인',
        '반복: 문제가 해결될 때까지 위 과정 반복',
    ])

    add_para(doc, '자주 발생하는 에러와 해결법:', bold=True)
    add_simple_table(doc,
        ['에러 유형', '원인', '해결 방법'],
        [
            ['FileNotFoundError', '파일 경로가 틀림', '.igs 또는 TK1.xlsx 파일 위치 확인'],
            ['ValueError: Sheet not found', 'Excel에 zero-1 시트 없음', 'TK1.xlsx에 zero-1 시트 존재 확인\n또는 환경변수 설정'],
            ['PermissionError', 'Excel 파일이 열려있음', 'Excel에서 TK1.xlsx 닫기 후 재실행'],
            ['ImportError', '필요한 라이브러리 미설치', 'pip install numpy openpyxl 등 실행'],
        ])

    doc.add_page_break()

    # ── 2D. 프로그램 실행방법 ──
    add_heading(doc, 'D. 프로그램 실행 방법', level=2)

    add_heading(doc, '① 방법 1: EXE 파일로 실행 (가장 간단)', level=3)
    add_para(doc, 'Python 설치 없이 실행파일만으로 사용하는 방법:', bold=True)
    add_numbered(doc, [
        '동일 폴더에 다음 3개 파일을 배치:\n'
        '    • SpringCalculator.exe (실행파일)\n'
        '    • TK1.xlsx (결과가 기록될 Excel 템플릿)\n'
        '    • 분석대상.igs (스프링 IGS 데이터 파일)',
        'Excel에서 TK1.xlsx가 열려있다면 반드시 닫기',
        'SpringCalculator.exe를 더블클릭하여 실행',
        '완료 후 TK1.xlsx의 zero-1 시트에서 결과 확인',
        '분석 이미지(spring_detailed_analysis.png 등)도 같은 폴더에 생성됨',
    ])

    add_heading(doc, '② 방법 2: Python 스크립트 직접 실행', level=3)
    add_para(doc, '개발/수정이 필요할 때는 Python을 설치하고 직접 실행:', bold=True)
    add_code(doc, '''
# 1단계: Python 가상환경 활성화
cd F:\\SpringPitch
.\\.venv\\Scripts\\Activate.ps1

# 2단계: 메인 분석 실행
python detailed_center_analysis.py

# 3단계: (선택) Excel 데이터 시각화
python plot_tk1_excel.py

# 4단계: (선택) 원시 좌표 CSV 내보내기
python tools\\export_xyz_and_plot.py
''')

    add_heading(doc, '③ 방법 3: 명령줄에서 IGS 파일 지정 실행', level=3)
    add_code(doc, '''
# 특정 IGS 파일을 인자로 지정
F:/SpringPitch/.venv/Scripts/python.exe detailed_center_analysis.py "C:\\데이터\\my_spring.igs"

# EXE 버전
SpringCalculator.exe "C:\\데이터\\my_spring.igs"
''')

    add_heading(doc, '④ 환경 변수 설정 (고급)', level=3)
    add_para(doc, '프로그램 동작을 세밀하게 조정할 수 있는 환경 변수:', italic=True)
    add_simple_table(doc,
        ['환경 변수', '기본값', '설명'],
        [
            ['SPRING_NORM_METHOD', 'linear', '표준화 방법 (linear / uniform_bspline / nurbs)'],
            ['SPRING_THETA_MODE', 'start0_unwrapped', '각도 처리 방식'],
            ['SPRING_THETA_POSITIVE', 'true', '양의 방향 증가 강제'],
            ['SPRING_START_STRATEGY', 'min_radius', '시작점 선택 전략'],
            ['SPRING_CREATE_ZERO1_IF_MISSING', 'false', 'zero-1 시트 자동 생성 여부'],
        ])

    doc.add_page_break()

    # ── 2E. 출력물 이해 ──
    add_heading(doc, 'E. 프로그램 출력물 이해', level=2)

    add_heading(doc, '① Excel 출력 (TK1.xlsx)', level=3)
    add_bullets(doc, [
        'raw 시트: 원시 IGS 좌표 전체 (3,141개)',
        'zero-1 시트: 표준화된 1,000개 포인트 + 분석 결과 (A~X 컬럼)',
        '기존 서식/포맷은 보존되며, 데이터만 업데이트됨',
    ])

    add_heading(doc, '② 이미지 출력', level=3)
    add_simple_table(doc,
        ['파일명', '내용'],
        [
            ['spring_detailed_analysis.png', '6개 서브플롯: 국소좌표, R/Z vs 각도, 상면도, 측면도, 3D, 비교'],
            ['new_result_report.png', '12개 서브플롯: 메트릭 요약, R/D/H/Pitch 프로파일, 분포, 상관행렬'],
        ])

    add_heading(doc, '③ CSV 출력', level=3)
    add_bullets(doc, [
        'output/points_xyz.csv: 원시 좌표 전체 (x, y, z)',
    ])

    doc.add_page_break()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 3. 기타
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_heading(doc, '3. 기타', level=1)

    # ── 3A. Q&A ──
    add_heading(doc, 'A. 자주 묻는 질문 (Q&A)', level=2)

    qa_pairs = [
        ('Q1: Python을 전혀 모르는데 프로그램을 실행할 수 있나요?',
         '네, EXE 파일을 사용하면 Python 설치 없이 실행할 수 있습니다. '
         'SpringCalculator.exe, TK1.xlsx, .igs 파일을 같은 폴더에 두고 '
         'exe를 더블클릭하면 됩니다.'),
        ('Q2: IGS 파일의 포인트 수가 다른 경우(1000개가 아닌 경우)에도 작동하나요?',
         '네, 프로그램은 원시 포인트 수에 관계없이 항상 1,000개로 표준화합니다. '
         '3,141개든 500개든 등간격으로 리샘플링됩니다.'),
        ('Q3: 결과가 이상해 보이면 어떻게 하나요?',
         '① spring_detailed_analysis.png 이미지를 확인하여 3D 형상이 정상인지 확인\n'
         '② IGS 파일이 올바른 형식인지 확인 (엔터티 116이 포함되어 있는지)\n'
         '③ TK1.xlsx가 기존 데이터로 오염되지 않았는지 확인\n'
         '④ 에러 메시지를 AI에 질문하여 원인 파악'),
        ('Q4: 코드를 수정하고 싶은데 어디를 바꿔야 하나요?',
         '주요 계산 로직은 detailed_center_analysis.py의 main() 함수에 있습니다. '
         'IGS 파싱 방식을 바꾸려면 iges_parser.py를, '
         'Excel 출력 형식을 바꾸려면 write_excel_tk1() 함수를 수정합니다.'),
        ('Q5: 다른 CAD 파일(STEP, STL 등)도 분석할 수 있나요?',
         '현재는 IGES(.igs) 포맷만 지원합니다. 다른 포맷은 먼저 IGES로 변환하거나, '
         '별도의 파서를 추가해야 합니다.'),
        ('Q6: 단위가 inch인 IGS 파일은?',
         '프로그램이 Global 섹션에서 단위를 자동 감지하고 mm로 변환합니다. '
         'inch 단위 파일도 정상 처리됩니다.'),
        ('Q7: EXE를 새로 빌드(생성)하려면?',
         'Python과 PyInstaller가 설치된 환경에서 아래 명령을 실행합니다:\n'
         'pyinstaller --onefile --name SpringCalculator detailed_center_analysis.py'),
        ('Q8: 스프링이 아닌 다른 형상에도 사용할 수 있나요?',
         '이 프로그램은 나선(스프링) 형상에 특화되어 있습니다. '
         'PCA로 주축을 찾고 원통좌표로 변환하는 방식이므로, '
         '나선형 구조가 아닌 형상에는 의미 있는 결과가 나오지 않습니다.'),
    ]

    for q, a in qa_pairs:
        p_q = doc.add_paragraph()
        run_q = p_q.add_run(q)
        run_q.bold = True
        run_q.font.size = Pt(11)
        p_a = doc.add_paragraph()
        p_a.paragraph_format.left_indent = Cm(0.5)
        run_a = p_a.add_run(a)
        run_a.font.size = Pt(10.5)
        doc.add_paragraph()  # 간격

    # ── 3B. 참고 자료 ──
    add_heading(doc, 'B. 참고 자료 및 용어 정리', level=2)

    add_heading(doc, '① 핵심 용어 사전', level=3)
    add_simple_table(doc,
        ['용어', '영문', '설명'],
        [
            ['IGS/IGES', 'Initial Graphics Exchange Specification', 'CAD 데이터 교환 표준 포맷'],
            ['PCA', 'Principal Component Analysis', '데이터의 주요 방향(주축)을 찾는 통계 방법'],
            ['리샘플링', 'Resampling', '데이터 점 수를 변경 (등간격으로 재배치)'],
            ['원통좌표계', 'Cylindrical Coordinates', '(R, θ, Z)로 위치를 표현하는 좌표계'],
            ['호 길이', 'Arc Length', '곡선을 따른 실제 거리'],
            ['피치', 'Pitch', '스프링 코일 1회전당 축방향 이동 거리'],
            ['직경 (수직)', 'Perpendicular Diameter', '2×R (축에서 수직 방향 지름)'],
            ['직경 (벡터)', 'Vectorial Diameter', '반 바퀴 앞 점까지의 직선 거리'],
            ['Turn', 'Turn (Revolution)', '누적 회전수 (360° = 1 turn)'],
            ['Unwrap', 'Phase Unwrapping', '각도 불연속(±180°)을 제거하여 연속값으로 변환'],
            ['국소좌표계', 'Local Coordinate System', '스프링 주축 기준으로 재정의한 좌표계'],
            ['엔터티', 'Entity', 'IGES 파일 내의 개별 기하학적 객체 (점, 선, 면 등)'],
        ])

    add_heading(doc, '② 참고 웹사이트', level=3)
    add_bullets(doc, [
        'Python 공식 사이트: https://www.python.org',
        'NumPy 문서: https://numpy.org/doc',
        'Matplotlib 문서: https://matplotlib.org',
        'IGES 표준 설명: https://wiki.eclipse.org/IGES',
        'VS Code 다운로드: https://code.visualstudio.com',
    ])

    add_heading(doc, '③ 교육 후 추천 학습 경로', level=3)
    add_numbered(doc, [
        '파이썬 기초 문법 온라인 강의 (점프 투 파이썬, 코드잇 등) - 1~2주',
        'NumPy 기초 (배열 연산, 인덱싱) - 3일',
        'Matplotlib 시각화 기초 (plot, scatter, subplot) - 2일',
        '본 프로젝트 코드를 한 함수씩 읽으며 AI에게 질문 - 1주',
        '간단한 수정(파라미터 변경) 시도 → 결과 확인 반복',
    ])

    doc.add_page_break()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 부록 1: iges_parser.py 상세 기능 설명
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_heading(doc, '부록 1: iges_parser.py — 함수/모듈 상세 기능', level=1)
    add_para(doc, 'SpringCalculator.exe를 구성하는 2개 파이썬 파일 중 하나로, '
             'IGES 파일을 섹션 단위로 파싱하여 3D 포인트 좌표를 추출하는 전문 파서입니다. '
             'detailed_center_analysis.py에서 import되어 사용됩니다.', italic=True)

    # ── 데이터 클래스 ──
    add_heading(doc, '① 데이터 클래스 (자료 구조)', level=2)
    add_para(doc, 'iges_parser.py는 @dataclass 데코레이터를 이용하여 '
             'IGES 파일의 핵심 정보를 담는 4개의 데이터 클래스를 정의합니다:')

    add_simple_table(doc,
        ['클래스명', '역할', '주요 필드'],
        [
            ['IgesGlobal', 'Global 섹션 메타데이터',
             'param_delim (구분자, 기본 ",")\nrecord_delim (끝표시, 기본 ";")\n'
             'units_name ("MM" 또는 "IN")\nunits_flag (2=mm, 1=inch)\n'
             'scale_to_mm (mm 변환 배율)'],
            ['Transform124', '좌표 변환행렬\n(엔터티 124)',
             'm: 4×4 float 행렬\n(3×4 회전+이동 + 마지막 행 [0,0,0,1])\n'
             'apply(x,y,z) 메서드로 좌표 변환'],
            ['Point116', '포인트 데이터\n(엔터티 116)',
             'x, y, z: 좌표값\n'
             'de_index: Directory Entry 인덱스\n'
             'xform_ptr: 변환행렬 포인터'],
            ['DirectoryEntry', 'Directory 섹션의\n엔터티 메타 정보',
             'entity_type: 엔터티 종류 (116, 124 등)\n'
             'param_pointer: P섹션 시작 위치\n'
             'xform_ptr: 변환행렬 참조\n'
             'param_line_count: P섹션 줄 수'],
        ])

    add_code(doc, '''
# 데이터 클래스 예시
@dataclass
class IgesGlobal:
    param_delim: str = ','      # 파라미터 구분자
    record_delim: str = ';'     # 레코드 끝 표시
    units_name: str = 'MM'      # 단위 이름
    units_flag: int = 2         # 단위 코드 (2=mm)
    scale_to_mm: float = 1.0    # mm 변환 배율

@dataclass
class Transform124:
    m: List[List[float]]        # 4×4 변환 행렬
    
    def apply(self, x, y, z):   # 좌표에 행렬 곱 적용
        vx = m[0][0]*x + m[0][1]*y + m[0][2]*z + m[0][3]
        ...
        return vx, vy, vz
''')

    # ── 내부 함수 ──
    add_heading(doc, '② 내부(private) 함수 — 파일 읽기 및 섹션 분리', level=2)
    add_para(doc, '밑줄(_)로 시작하는 함수는 모듈 내부에서만 쓰이는 보조 함수입니다:')

    add_simple_table(doc,
        ['함수명', '기능 설명', '입력 → 출력'],
        [
            ['_chunks_80(filepath)',
             'IGES 파일을 한 줄씩 읽어\n73열의 섹션 문자(S/G/D/P/T)로\n5개 리스트에 분류',
             '파일 경로 →\n(S줄들, G줄들, D줄들, P줄들, T줄들)'],
            ['_parse_global_delimiters(G_lines)',
             'G 섹션에서 파라미터 구분자와\n레코드 구분자를 추출\n(보통 "," 와 ";")',
             'G 섹션 줄 리스트 →\n(param_delim, record_delim)'],
            ['_parse_global_units(G_lines)',
             'G 섹션에서 단위 정보 추출\n"2HMM" 패턴으로 밀리미터 감지',
             'G 섹션 줄 리스트 →\n(units_name, units_flag)'],
            ['_unit_scale_to_mm(name, flag)',
             '단위 이름/코드를 mm 배율로 변환\nMM→1.0, INCH→25.4',
             '단위 정보 → float 배율'],
            ['_parse_directory(D_lines)',
             'D 섹션의 80자×2줄 레코드를\n파싱하여 DirectoryEntry 리스트 생성\n(8자 필드 20개 분해)',
             'D 섹션 줄 리스트 →\nList[DirectoryEntry]'],
            ['_tokenize_p_records(P_lines)',
             'P 섹션의 각 줄을 시퀀스 번호별로\n맵핑 (여러 줄에 걸친 데이터 병합)',
             'P 섹션 줄 리스트 →\nDict[seq_num → param_text]'],
            ['_float_token(token)',
             '문자열 토큰을 float로 변환\nIGES 특유의 "D" 지수 표기도 처리\n예: "1.23D+02" → 123.0',
             '문자열 → float'],
        ])

    add_tip_box(doc, '언더스코어(_) 함수명 규칙',
                '파이썬에서 함수명이 _로 시작하면 "이 함수는 외부에서 직접 호출하지 말라"는 '
                '관례적 표시입니다. iges_parser.py에서는 parse_iges_points()만 외부에서 호출하며, '
                '나머지 _함수들은 내부 처리 단계입니다.')

    # ── 핵심 공개 함수 ──
    add_heading(doc, '③ 핵심 공개 함수: parse_iges_points()', level=2)
    add_para(doc, '외부에서 호출하는 유일한 진입 함수입니다. '
             'IGS 파일 경로를 받아 모든 포인트의 (X, Y, Z) 좌표 리스트를 반환합니다.', bold=True)

    add_code(doc, '''
def parse_iges_points(filepath: str) -> Tuple[List[Tuple[float,float,float]], IgesGlobal]:
    """IGES 파일을 파싱하여 모든 포인트(116) 좌표를 추출"""
    
    # 1단계: 파일 → 5개 섹션(S/G/D/P/T)으로 분리
    S, G, D, P, T = _chunks_80(filepath)
    
    # 2단계: Global 섹션 → 구분자, 단위, 스케일 결정
    param_delim, rec_delim = _parse_global_delimiters(G)
    units_name, units_flag = _parse_global_units(G)
    scale = _unit_scale_to_mm(units_name, units_flag)
    
    # 3단계: Directory 섹션 → 엔터티 목록 생성
    entries = _parse_directory(D)
    
    # 4단계: Parameter 섹션 → 시퀀스별 파라미터 맵
    pmap = _tokenize_p_records(P)
    
    # 5단계: 변환행렬(124) 수집 → 4×4 행렬 딕셔너리
    transforms = {}
    for de in entries:
        if de.entity_type == 124:
            ... # P 섹션에서 12개 값 → 4×4 행렬
    
    # 6단계: 포인트(116) 추출 + 변환행렬 적용 + mm 스케일링
    points = []
    for de in entries:
        if de.entity_type == 116:
            x, y, z = ... # P 섹션에서 좌표 읽기
            if de.xform_ptr in transforms:
                x, y, z = transforms[de.xform_ptr].apply(x, y, z)
            points.append((x * scale, y * scale, z * scale))
    
    return points, glb   # 좌표 리스트 + 글로벌 메타데이터
''')

    add_para(doc, '처리 흐름 요약:', bold=True)
    add_code(doc, '''
┌────────────┐   ┌──────────────┐   ┌──────────────┐   ┌────────────┐
│ _chunks_80 │──▶│_parse_global │──▶│_parse_direc- │──▶│_tokenize_  │
│ (파일 분리)  │   │ (구분자/단위)  │   │ tory (DE목록) │   │ p_records   │
└────────────┘   └──────────────┘   └──────────────┘   └──────┬─────┘
                                                              │
     ┌────────────────────────────────────────────────────────┘
     ▼
┌──────────────────────────────────────────┐
│  parse_iges_points (핵심 로직)             │
│  • 124 변환행렬 수집                       │
│  • 116 포인트 좌표 추출                    │
│  • 변환행렬 적용 + mm 단위 스케일링         │
│  → [(x1,y1,z1), (x2,y2,z2), ...] 반환    │
└──────────────────────────────────────────┘
''')

    doc.add_page_break()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 부록 2: detailed_center_analysis.py 상세 기능 설명
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_heading(doc, '부록 2: detailed_center_analysis.py — 함수/모듈 상세 기능', level=1)
    add_para(doc, 'SpringCalculator.exe의 메인 엔트리포인트 파일입니다. '
             'IGS 로드부터 좌표 변환, 분석, Excel 기록, 시각화까지 전체 파이프라인이 '
             '이 한 파일의 34개 함수와 main() 함수에 구현되어 있습니다.\n\n'
             '파일 크기: 약 1,400행 | 함수: 34개 | import: numpy, matplotlib, openpyxl, pandas 등',
             italic=True)

    # ── 환경변수/설정 상수 ──
    add_heading(doc, '① 환경변수 기반 설정 상수 (30~70행)', level=2)
    add_para(doc, '프로그램 동작을 외부에서 조정할 수 있도록 '
             'os.getenv()로 환경변수를 읽어 전역 상수로 설정합니다:')

    add_simple_table(doc,
        ['상수명', '환경변수', '기본값', '역할'],
        [
            ['NORMALIZATION_METHOD', 'SPRING_NORM_METHOD', 'linear', '표준화 방법 (linear/bspline/nurbs)'],
            ['THETA_MODE', 'SPRING_THETA_MODE', 'start0_unwrapped', 'θ 각도 처리 모드'],
            ['THETA_POSITIVE', 'SPRING_THETA_POSITIVE', 'true', 'θ 양의 방향 강제'],
            ['START_STRATEGY', 'SPRING_START_STRATEGY', 'min_radius', '시작점 선택 전략'],
            ['SIMPLE_CONVERT', 'SPRING_SIMPLE_CONVERT', 'true', '단순 원통 변환 모드'],
            ['ORIGIN_MODE', 'SPRING_ORIGIN_MODE', 'start', '원점 설정 방식 (start/midpoint)'],
            ['OUTLIER_ENABLE', 'SPRING_OUTLIER_ENABLE', 'true', '이상치 감지 활성화'],
            ['VECTORIAL_TURN_OFFSET', 'SPRING_VECTORIAL_TURN_OFFSET', '0.5', '벡터직경 반바퀴 오프셋'],
        ])

    add_tip_box(doc, '환경변수란?',
                '운영체제가 프로그램에 전달하는 "이름=값" 설정입니다.\n'
                'PowerShell에서 실행 전에  $env:SPRING_NORM_METHOD = "uniform_bspline"  처럼 설정하면\n'
                '프로그램이 그 값을 읽어 동작을 변경합니다. 설정하지 않으면 기본값이 쓰입니다.')

    # ── 함수 그룹 A: 데이터 입력 ──
    add_heading(doc, '② 데이터 입력 함수', level=2)

    add_simple_table(doc,
        ['함수명', '행번호', '기능 설명', '호출 관계'],
        [
            ['parse_igs_points(filepath)',
             '73행',
             'IGS 파일에서 좌표 추출\n'
             '① iges_parser 모듈 호출 시도\n'
             '② 실패 시 정규식 폴백으로 재시도\n'
             '③ numpy 배열로 변환하여 반환',
             'main()에서 호출\n→ iges_parser.parse_iges_points()\n   호출 (try/except)'],
        ])

    add_code(doc, '''
# parse_igs_points() - 이중 파싱 전략
def parse_igs_points(filepath):
    try:
        # 방법 1: 전문 섹션 파서 (iges_parser.py)
        from iges_parser import parse_iges_points as _parse_iges
        pts, glb = _parse_iges(filepath)
        return np.array(pts, dtype=float)
    except Exception:
        # 방법 2: 정규식으로 간이 파싱 (폴백)
        pattern = r"116,\\s*(...),\\s*(...),\\s*(...)...;"
        # 파일을 줄 단위로 읽으며 패턴 매칭
        ...
''')

    # ── 함수 그룹 B: 축/좌표 분석 ──
    add_heading(doc, '③ 축 분석 및 좌표 변환 함수', level=2)
    add_para(doc, '스프링의 주축을 찾고, 전역 좌표를 국소 좌표 → 원통 좌표로 변환하는 핵심 함수들:')

    add_simple_table(doc,
        ['함수명', '행', '기능 설명'],
        [
            ['find_spring_axis_pca(points)',
             '97',
             'PCA(주성분분석)으로 스프링 주축 방향 계산\n'
             '• 입력: 3D 좌표 배열 (N×3)\n'
             '• 공분산 행렬 → 고유값 분해\n'
             '• 출력: 무게중심(center), 주축방향(axis), 고유값, 고유벡터'],
            ['build_local_frame(axis_direction)',
             '119',
             '주축 방향으로부터 직교 좌표계 (Ex, Ey, Ez) 구성\n'
             '• Ez = 축 방향 (정규화)\n'
             '• Ex = 참조벡터와 Ez의 외적\n'
             '• Ey = Ez × Ex\n'
             '• 출력: 회전행렬 R + 3개 기저벡터'],
            ['to_local_coordinates(points, origin, axis)',
             '140',
             '글로벌 좌표 → 국소 직교좌표 변환\n'
             '• 원점(origin) 기준으로 이동\n'
             '• 주축 = 로컬 X축으로 정렬\n'
             '• 출력: (x_local, y_local, z_local) 배열'],
            ['cylindrical_from_local(local_xyz)',
             '181',
             '국소 직교좌표 → 원통좌표 (R, θ, Z) 변환\n'
             '• R = √(x² + y²), θ = atan2(y, x)\n'
             '• Z = z (축 방향 그대로)'],
            ['unwrap_theta(theta, sort_idx)',
             '190',
             'θ의 ±π 불연속을 제거하여 연속값으로 변환\n'
             '• numpy.unwrap() 활용\n'
             '• sort_idx 제공 시 해당 순서로 정렬 후 처리'],
        ])

    # ── 함수 그룹 C: 정렬/표준화 ──
    add_heading(doc, '④ 포인트 정렬 및 표준화 함수', level=2)

    add_simple_table(doc,
        ['함수명', '행', '기능 설명'],
        [
            ['order_points_by_helical_parameter(...)',
             '713',
             '나선(헬리컬) 파라미터 기반으로 포인트 순서 결정\n'
             '• 축 투영만으로는 코일 끝단에서 순서 오류 발생\n'
             '• YZ 평면의 각도(unwrapped)를 기준으로 정렬\n'
             '• 축 진행 방향과의 상관 관계로 방향 보정'],
            ['resample_curve(points, num_points=1000)',
             '734',
             '3D 곡선을 등간격으로 리샘플링\n'
             '• 누적 호 길이(arc length) 계산\n'
             '• 총 길이를 num_points 구간으로 등분\n'
             '• 각 축(x,y,z)별 np.interp()으로 보간\n'
             '• 출력: 1,000개의 등간격 좌표 배열'],
            ['sort_points_along_axis(points, axis_pt, axis_dir)',
             '707',
             '축 투영 기반 단순 정렬 (보조용)\n'
             '• 주축 방향 투영 → argsort'],
            ['project_points_to_axis(points, axis_pt, axis_dir)',
             '279',
             '각 점을 주축에 투영\n'
             '• 축 상의 투영 길이, 축으로부터 거리, 축 위의 점 반환'],
            ['_determine_origin_and_startidx(local_coords)',
             '755',
             '시작 인덱스 결정 (현재 항상 0 반환)'],
        ])

    # ── 함수 그룹 D: 치수 계산 ──
    add_heading(doc, '⑤ 치수 계산 함수', level=2)

    add_simple_table(doc,
        ['함수명', '행', '기능 설명'],
        [
            ['compute_basic_metrics(local_xyz, r, θ, z)',
             '215',
             'Step 4 기초 치수 산출\n'
             '• Turn = θ_unwrapped / 2π\n'
             '• Pitch = dZ / dTurn (차분)\n'
             '• Diameter = 2 × R (기초값)\n'
             '• 이동평균으로 평활화\n'
             '• 출력: 컬럼 딕셔너리 + 요약 통계'],
            ['moving_average(x, window)',
             '202',
             '단순 이동평균 필터\n'
             '• 엣지는 패딩(edge)으로 처리\n'
             '• R, Z, Pitch 노이즈 제거에 사용'],
            ['nine_if_small(n)',
             '272',
             '데이터 길이에 따른 이동평균 윈도우 자동 결정\n'
             '• 최소 5, 최대 101, 홀수 보장'],
            ['compute_curvature(points)',
             '769',
             '3D 곡선의 곡률(κ) 계산\n'
             '• κ = |P′ × P″| / |P′|³\n'
             '• 이상치(outlier) 감지 기준으로 활용'],
            ['theta_modes(local_std, start_index)',
             '808',
             '다양한 θ 모드 계산\n'
             '• raw(원시), unwrapped(연속), start0(시작점 0°)\n'
             '• 딕셔너리로 반환'],
            ['wrap_to_pi(x)',
             '831',
             '각도를 [-π, π] 범위로 래핑'],
            ['enforce_positive(theta_rad)',
             '835',
             'θ가 양의 방향으로 증가하도록 부호 보정'],
        ])

    # ── 함수 그룹 E: 스프링 구조 분석 ──
    add_heading(doc, '⑥ 스프링 구조 분석 함수', level=2)

    add_simple_table(doc,
        ['함수명', '행', '기능 설명'],
        [
            ['analyze_spring_layers(points, axis_pt, axis_dir, n=20)',
             '298',
             '스프링을 축 방향으로 n개 층으로 분할하여\n'
             '각 층의 중심점, 평균 반지름, Z값 산출'],
            ['find_optimal_spring_center(points)',
             '330',
             'PCA 축 + 층별 분석을 종합하여 최적 중심축 결정\n'
             '• 상위 10% 포인트에서 상단 중심 계산\n'
             '• 상위 3개 층 중심 평균으로 정밀화\n'
             '• 딕셔너리로 pca_center, top_center 등 반환'],
            ['_select_start_end_by_strategy(...)',
             '696',
             '시작/끝 구간을 축 투영 기반으로 선택\n'
             '• 축 방향 양단 end_slice_percent% 마스크 생성'],
        ])

    # ── 함수 그룹 F: 출력/저장 ──
    add_heading(doc, '⑦ 출력 및 저장 함수', level=2)

    add_simple_table(doc,
        ['함수명', '행', '기능 설명'],
        [
            ['write_excel_tk1(path, sheet, std_points, ...)',
             '420',
             'TK1.xlsx의 zero-1 시트에 분석 결과 기록\n'
             '• 기존 서식 보존, 대상 컬럼만 업데이트\n'
             '• 컬럼: No, x,y,z, x_norm,y_norm,z_norm, R,θ,Z\n'
             '• 보조 컬럼 L~X (Turn, Diameter, Pitch 등)\n'
             '• 요약 블록(Z3:AC14) 기록 기능\n'
             '• L~X 헤더 하드코딩으로 일관성 보장'],
            ['write_excel_raw(path, sheet, raw_points)',
             '557',
             'raw 시트에 원시 좌표 전체 기록\n'
             '• 기존 데이터 클리어 후 재기록\n'
             '• 컬럼: No, x, y, z'],
            ['safe_save_workbook(wb, path, logger)',
             '378',
             '안전한 Excel 저장\n'
             '• 임시 파일에 먼저 저장 → os.replace()로 교체\n'
             '• 저장 중 충돌/손상 방지'],
            ['write_csv(path, header, columns)',
             '409',
             'CSV 파일로 데이터 내보내기\n'
             '• 헤더 + 딕셔너리 기반 컬럼 데이터 기록'],
            ['ensure_dir(path)',
             '374',
             '디렉토리 존재 확인 및 생성 (os.makedirs)'],
        ])

    # ── 함수 그룹 G: 시각화 ──
    add_heading(doc, '⑧ 시각화 함수', level=2)

    add_simple_table(doc,
        ['함수명', '행', '기능 설명'],
        [
            ['visualize_spring_analysis(result, ...)',
             '585',
             '종합 분석 시각화 (3×2 서브플롯)\n'
             '• [1] 국소좌표 vs 인덱스\n'
             '• [2] R/Z vs 언래핑 각도\n'
             '• [3] 상면도 (Y vs X)\n'
             '• [4] 측면도 (Z vs X)\n'
             '• [5] 3D 리샘플 + 이상치 마커\n'
             '• [6] zero-1 vs std 비교\n'
             '→ spring_detailed_analysis.png 저장'],
            ['create_new_result_report_chart(excel_path, out)',
             '849',
             '결과 보고서 차트 (3×4 = 12개 서브플롯)\n'
             '• 메트릭 요약, R/Diam/Height 프로파일\n'
             '• Pitch 분포, θ vs Turn, 히스토그램\n'
             '• 3D 산점도, 상관행렬, 품질지표\n'
             '→ new_result_report.png 저장'],
            ['plot_cylindrical_3d(points, r, θ, z, path)',
             '799',
             '원통좌표 3D 플롯 (보조 함수)'],
            ['read_std_sheet(filepath)',
             '684',
             'Excel에서 std 시트를 pandas로 읽기\n'
             '(비교 시각화용)'],
        ])

    # ── main() 함수 ──
    add_heading(doc, '⑨ main() 함수 — 전체 파이프라인 실행', level=2)
    add_para(doc, 'EXE 또는 python 실행 시 호출되는 메인 함수(1052행~끝)입니다. '
             '전체 분석 파이프라인을 순서대로 실행합니다:', bold=True)

    add_code(doc, '''
def main():
    # ── 1. 초기 설정 ──
    # 로깅 설정, 실행 디렉토리 결정
    # IGS 파일 탐색 (인자 → 실행폴더 → 부모폴더 → CWD)
    # TK1.xlsx 위치 확인
    
    # ── 2. 기존 데이터 초기화 ──
    clear_zero1_data()  # zero-1 시트 A:X 데이터 삭제 (서식 보존)
    
    # ── 3. Step 1: IGS 파싱 ──
    original_points = parse_igs_points(filepath)  # → 3,141개 좌표
    
    # ── 4. PCA 축 결정 ──
    center, axis_vector, _, _ = find_spring_axis_pca(original_points)
    
    # ── 5. Step 2: 정렬 + 리샘플링 ──
    points_sorted, _, _ = order_points_by_helical_parameter(...)
    resampled_points = resample_curve(points_sorted, num_points=1000)
    local_coords_step2, _ = to_local_coordinates(resampled_points, ...)
    # 방향 보정, X시작=0, YZ 중심=0 처리
    
    # ── 6. Step 3: 원통좌표 변환 ──
    cyl_Z = local_coords_step2[:, 0]   # x_norm → Z
    r_step3 = sqrt(y² + z²)             # → R
    theta = atan2(z, y) → unwrap → 시작=0°  # → θ
    
    # ── 7. Step 4: Excel 수식 계산 (L~X 컬럼) ──
    col_N = theta_unwrapped / 360        # Turn
    col_Q = R[i] + R[half_turn_back]     # Perpendicular Diameter
    col_S = distance(raw[i], raw[half_turn_back])  # Vectorial Diameter
    col_W = Z[i] - Z[full_turn_back]     # Pitch
    col_X = min(W where W > 0)           # min_Pitch
    
    # ── 8. 출력 ──
    write_excel_raw(...)                  # raw 시트
    write_excel_tk1(...)                  # zero-1 시트 (전체 결과)
    
    # ── 9. 시각화 ──
    visualize_spring_analysis(...)        # 분석 6-plot
    create_new_result_report_chart(...)   # 보고서 12-plot
''')

    add_heading(doc, '⑩ main() 파일 탐색 우선순위', level=2)
    add_para(doc, 'main()은 IGS 파일과 TK1.xlsx를 아래 순서로 자동 탐색합니다:')
    add_simple_table(doc,
        ['우선순위', 'IGS 파일 탐색', 'TK1.xlsx 탐색'],
        [
            ['1순위', '명령줄 인자 (sys.argv[1])', 'IGS 파일과 동일 폴더'],
            ['2순위', '실행 폴더의 첫 번째 .igs', '스크립트(EXE) 폴더'],
            ['3순위', '부모 폴더의 .igs', '부모 폴더'],
            ['4순위', '현재 작업 디렉토리(CWD)', '현재 작업 디렉토리'],
        ])

    add_heading(doc, '⑪ 두 파일의 관계 요약', level=2)
    add_code(doc, '''
┌──────────────────────────────────────────────────────────────┐
│              SpringCalculator.exe 구성                        │
├──────────────────────────────────────────────────────────────┤
│                                                              │
│  detailed_center_analysis.py  (메인 엔트리포인트)              │
│  ├── main()                   프로그램 시작점                  │
│  ├── parse_igs_points()       IGS 파싱 (→ iges_parser 호출)   │
│  ├── find_spring_axis_pca()   PCA 축 분석                     │
│  ├── to_local_coordinates()   좌표 변환                       │
│  ├── resample_curve()         1000개 리샘플링                  │
│  ├── cylindrical_from_local() R,θ,Z 변환                     │
│  ├── write_excel_tk1()        Excel 결과 기록                 │
│  ├── visualize_spring_analysis()  분석 시각화                  │
│  └── create_new_result_report_chart()  보고서 차트             │
│          │                                                    │
│          │  from iges_parser import parse_iges_points          │
│          ▼                                                    │
│  iges_parser.py  (IGES 전문 파서)                              │
│  ├── parse_iges_points()      유일한 공개 API                  │
│  ├── _chunks_80()             파일 섹션 분리                   │
│  ├── _parse_global_*()        메타데이터 파싱                  │
│  ├── _parse_directory()       DE 목록 생성                    │
│  ├── _tokenize_p_records()    파라미터 맵핑                   │
│  └── _float_token()           숫자 변환 (D/E 지수)            │
│                                                              │
│  + PyInstaller 번들: numpy, matplotlib, openpyxl, pandas 등  │
└──────────────────────────────────────────────────────────────┘
''')

    # 저장
    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)
    print(f'교육자료 생성 완료: {OUT_PATH}')
    return OUT_PATH


if __name__ == '__main__':
    build_training_doc()
